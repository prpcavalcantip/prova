import os
import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from io import BytesIO
import re

# --- Configuração do Tema ---
st.set_page_config(page_title="Adaptador de Provas", layout="wide")

# --- Estilização CSS ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@400;700&display=swap');
    .main-title {
        font-family: 'Roboto', sans-serif;
        font-size: 2.5rem;
        color: #1E3A8A;
        font-weight: 700;
        margin-bottom: 0.5rem;
    }
    .subtitle {
        font-family: 'Roboto', sans-serif;
        font-size: 1rem;
        color: #64748B;
    }
    .stButton > button {
        background: linear-gradient(90deg, #FBBF24, #F59E0B);
        color: #1E3A8A;
        border-radius: 8px;
        padding: 0.5rem 1.5rem;
        font-family: 'Roboto', sans-serif;
        font-size: 1rem;
        border: none;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
        transition: transform 0.2s;
    }
    .stButton > button:hover {
        transform: scale(1.05);
        background: linear-gradient(90deg, #F59E0B, #FBBF24);
    }
    .stTextInput > div > input {
        border-radius: 8px;
        font-family: 'Roboto', sans-serif;
        border: 2px solid #1E3A8A;
    }
    .stSelectbox > div > select {
        border-radius: 8px;
        font-family: 'Roboto', sans-serif;
        border: 2px solid #1E3A8A;
    }
    .stFileUploader > div {
        border-radius: 8px;
        border: 2px dashed #1E3A8A;
    }
    .info-box {
        background: linear-gradient(135deg, #F3F4F6, #FFFFFF);
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #E2E8F0;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
    }
    .footer {
        text-align: center;
        color: #64748B;
        font-family: 'Roboto', sans-serif;
        margin-top: 2rem;
        font-size: 0.9rem;
    }
    </style>
""", unsafe_allow_html=True)

# --- Dados e Funções ---
PRIMARY_COLOR = "#1E3A8A"
SECONDARY_COLOR = "#FBBF24"

TIPOS_DICAS = {
    "TDAH": [
        "Leia com calma. Destaque palavras importantes.",
        "Use setas ou cores para conectar ideias.",
        "Evite distrações: foque uma questão de cada vez.",
        "Grife os conceitos-chave no enunciado.",
        "Relacione a questão com exemplos práticos."
    ],
    "Dislexia": [
        "Leia devagar. Palavras difíceis podem confundir.",
        "Separe frases longas em partes curtas.",
        "Use régua ou dedo para acompanhar.",
        "Leia em voz baixa para entender melhor.",
        "Marque palavras que aparecem muito."
    ],
    "TEA": [
        "Observe a estrutura lógica da questão.",
        "Use imagens mentais para entender conceitos.",
        "Evite interpretações subjetivas: vá direto ao que é pedido.",
        "Releia com atenção cada alternativa.",
        "Destaque padrões e repetições."
    ]
}

def obter_dica(perfil, numero_questao):
    dicas = TIPOS_DICAS.get(perfil, [])
    if 0 < numero_questao <= len(dicas):
        return dicas[numero_questao - 1]
    return "Leia com atenção."

def avaliar_enunciado(enunciado):
    palavras_complexas = ['explique', 'analise', 'interprete', 'justifique', 'discuta']
    score = len(enunciado)
    for termo in palavras_complexas:
        if termo in enunciado.lower():
            score += 20
    return score

def limpar_enunciado(enunciado):
    enunciado = re.sub(r'^QUESTÃO\s*\d+\s*', '', enunciado, flags=re.IGNORECASE)
    enunciado = re.sub(r'\s+', ' ', enunciado)
    return enunciado.strip()

def simplificar_enunciado_chatgpt(enunciado):
    enunciado = limpar_enunciado(enunciado)
    enunciado = re.sub(r'[Ss]egundo o texto,? ?', '', enunciado)
    enunciado = re.sub(r'[Dd]e acordo com o autor,? ?', '', enunciado)
    enunciado = re.sub(r'\s+', ' ', enunciado)
    return enunciado.strip()

def simplificar_questao(texto, perfil):
    texto = re.sub(r'\s+', ' ', texto).strip()
    match = re.match(r'(QUESTÃO \d+ )(.*?)([A-E]\))', texto)
    if not match:
        enunciado = texto
        alternativas = []
    else:
        inicio = match.start(3)
        enunciado = texto[:inicio].strip()
        alternativas = re.findall(r'([A-E]\))\s?(.*?)(?=\s[A-E]\)|$)', texto[inicio:])

    enunciado_simplificado = simplificar_enunciado_chatgpt(enunciado)
    return {
        "enunciado": enunciado_simplificado,
        "alternativas": alternativas,
        "score": avaliar_enunciado(enunciado_simplificado)
    }

def gerar_docx(questoes, nome_professor, materia, perfil):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Roboto'
    font.size = Pt(14)

    doc.add_heading(f"Prova Adaptada – {materia} – {nome_professor}", 0)
    doc.add_paragraph(f"Professor(a): {nome_professor}", style='Normal')
    doc.add_paragraph(f"Matéria: {materia}", style='Normal')
    doc.add_paragraph("")

    for i, q in enumerate(questoes, start=1):
        doc.add_paragraph(f"QUESTÃO {i}", style='Normal')
        doc.add_paragraph(q['enunciado'], style='Normal')
        doc.add_paragraph("")  # Espaço antes das alternativas

        for letra, alt in q['alternativas']:
            doc.add_paragraph(f"{letra} {alt.strip()}", style='Normal')

        doc.add_paragraph(f"🧠 DICA: {obter_dica(perfil, i)}", style='Normal')
        doc.add_paragraph("")

    return doc

def reset_form():
    st.session_state.pop("arquivo", None)
    st.session_state.pop("perfil", None)
    st.session_state.pop("prova_processada", None)
    st.session_state.pop("download_ready", None)

# --- Interface Principal ---
def main():
    # Sidebar
    with st.sidebar:
        st.markdown("<h3 style='font-family: Roboto; color: #1E3A8A;'>Sobre o Adaptador</h3>", unsafe_allow_html=True)
        st.markdown("""
            Ferramenta para adaptar provas em PDF para alunos neurodivergentes, com enunciados simplificados e dicas personalizadas.
            <br><br>
            <a href='https://colegioexodo.com' style='color: #FBBF24;'>Saiba mais</a>
        """, unsafe_allow_html=True)

    # Cabeçalho
    st.markdown("""
        <div style='text-align: center; margin-bottom: 2rem;'>
            <h1 class='main-title'>Adaptador de Provas Colégio Êxodo</h1>
            <p class='subtitle'>Ferramenta para professores – Versão Beta</p>
        </div>
    """, unsafe_allow_html=True)

    # Formulário
    with st.container():
        st.markdown("<div class='info-box'>", unsafe_allow_html=True)
        st.markdown("### 📋 Informações do Professor")
        col1, col2 = st.columns([3, 1])
        with col1:
            nome_professor = st.text_input("👤 Nome do Professor(a)", placeholder="Digite seu nome", help="Insira o nome do professor responsável pela prova.")
            materia = st.text_input("📘 Matéria", placeholder="Ex.: Matemática", help="Insira o nome da matéria da prova.")
        with col2:
            st.empty()  # Substitui st.image(LOGO_URL, width=100) por um espaço vazio
        st.markdown("</div>", unsafe_allow_html=True)

    if not nome_professor or not materia:
        st.info("📌 Preencha o nome do professor e a matéria para continuar.")
        st.stop()

    st.divider()

    if "prova_processada" not in st.session_state:
        st.markdown("### 📄 Configurar Adaptação")
        perfil = st.selectbox("💡 Tipo de adaptação:", list(TIPOS_DICAS.keys()), help="Escolha o perfil neurodivergente para adaptar a prova.", key="perfil")
        arquivo = st.file_uploader("📤 Envie a prova em PDF", type=["pdf"], key="arquivo", help="Selecione um arquivo PDF contendo a prova.")

        if arquivo:
            with st.spinner("⏳ Processando o PDF..."):
                try:
                    pdf = fitz.open(stream=arquivo.read(), filetype="pdf")
                except Exception as e:
                    st.error(f"❌ Erro ao abrir PDF: {e}")
                    return

                texto_completo = ""
                for pagina in pdf:
                    texto_completo += pagina.get_text()

                questoes_raw = re.findall(r'QUESTÃO \d+.*?(?=QUESTÃO \d+|$)', texto_completo, flags=re.DOTALL)
                if not questoes_raw:
                    st.error("❌ Nenhuma questão encontrada no PDF.")
                    return

                questoes_avaliadas = []
                for q_texto in questoes_raw:
                    questoes_avaliadas.append(simplificar_questao(q_texto, perfil))

                questoes_mais_diretas = sorted(questoes_avaliadas, key=lambda q: q['score'])[:5]

                doc = gerar_docx(questoes_mais_diretas, nome_professor, materia, perfil)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.session_state["prova_processada"] = buffer
                st.session_state["download_ready"] = True

    if st.session_state.get("download_ready", False):
        st.success("🎉 Prova adaptada gerada com sucesso!")
        st.markdown("""
            <div class='info-box'>
                <b>Prova pronta para download!</b> Faça o download do arquivo adaptado ou clique abaixo para adaptar outra prova.
            </div>
        """, unsafe_allow_html=True)

        st.download_button(
            label="📥 Baixar Prova Adaptada (.docx)",
            data=st.session_state["prova_processada"],
            file_name="prova_adaptada.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help="Clique para baixar a prova adaptada em formato Word."
        )

        st.divider()
        st.button("🔄 Adaptar outra prova", on_click=reset_form, help="Clique para reiniciar e adaptar uma nova prova.")

    # Rodapé
    st.markdown("<div class='footer'>Desenvolvido com ❤️ por Paulo Cavalcanti Pereira | Colégio Êxodo © 2025</div>", unsafe_allow_html=True)

if __name__ == "__main__":
    main()
